from pathlib import Path
from typing import Dict, Any
from .base import BaseConverter
from docx import Document

class DocxConverter(BaseConverter):
    SUPPORTED_EXTENSIONS = ['.docx']
    
    async def convert_to_markdown(self, file_path: Path) -> Dict[str, Any]:
        """Word转Markdown"""
        try:
            doc = Document(str(file_path))
            
            markdown_content = []
            current_paragraph = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    # 处理标题
                    style_name = paragraph.style.name if paragraph.style else ''
                    if style_name.startswith('Heading'):
                        try:
                            level = int(style_name[-1])
                        except (ValueError, IndexError):
                            level = 1
                        current_paragraph.append(f"{'#' * level} {paragraph.text}")
                        current_paragraph.append("")
                    else:
                        # 处理普通段落
                        current_paragraph.append(paragraph.text)
                        current_paragraph.append("")
            
            # 处理表格
            for table in doc.tables:
                table_content = self._convert_table_to_markdown(table)
                current_paragraph.append(table_content)
                current_paragraph.append("")
            
            # 处理图片（简化版）
            for rel in doc.part.rels.values():
                if "image" in rel.reltype:
                    current_paragraph.append(f"![图片](image_placeholder)")
                    current_paragraph.append("")
            
            return {
                "success": True,
                "content": "\n".join(current_paragraph),
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables)
            }
        except Exception as e:
            return {"success": False, "error": str(e)}
    
    def _convert_table_to_markdown(self, table) -> str:
        """将Word表格转换为Markdown表格"""
        rows = []
        
        # 处理表头
        if len(table.rows) > 0:
            header = " | ".join(cell.text for cell in table.rows[0].cells)
            rows.append(f"| {header} |")
            rows.append("| " + " | ".join(["---"] * len(table.rows[0].cells)) + " |")
        
        # 处理数据行
        for i, row in enumerate(table.rows):
            if i == 0:  # 跳过表头
                continue
            row_content = " | ".join(cell.text for cell in row.cells)
            rows.append(f"| {row_content} |")
        
        return "\n".join(rows)
    
    def _set_font(self, run_or_style_font, font_name='微软雅黑', size=None):
        """统一设置字体（含东亚字体，确保中文生效）"""
        from docx.oxml.ns import qn
        run_or_style_font.name = font_name
        # 设置东亚字体，这是中文生效的关键
        r = run_or_style_font._element
        rFonts = r.find(qn('w:rFonts'))
        if rFonts is None:
            from lxml import etree
            rFonts = etree.SubElement(r, qn('w:rFonts'))
        rFonts.set(qn('w:eastAsia'), font_name)
        if size:
            run_or_style_font.size = size

    async def convert_from_markdown(self, md_content: str, output_path: Path) -> Path:
        """Markdown转Word"""
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        style = doc.styles['Normal']
        self._set_font(style.font, '微软雅黑', Pt(11))

        # 统一设置内置标题样式的字体和字号
        heading_sizes = {0: 22, 1: 18, 2: 15, 3: 13, 4: 12}
        for i in range(5):
            try:
                h_style = doc.styles[f'Heading {i+1}'] if i > 0 else doc.styles['Title']
                self._set_font(h_style.font, '微软雅黑', Pt(heading_sizes.get(i, 12)))
                h_style.font.bold = True
                h_style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            except KeyError:
                pass

        lines = md_content.split('\n')
        in_code_block = False
        code_content = []

        for line in lines:
            line = line.strip()

            # 处理代码块
            if line.startswith('```'):
                if in_code_block:
                    # 结束代码块
                    p = doc.add_paragraph()
                    run = p.add_run('\n'.join(code_content))
                    self._set_font(run.font, 'Consolas', Pt(10))
                    code_content = []
                    in_code_block = False
                else:
                    in_code_block = True
                continue

            if in_code_block:
                code_content.append(line)
                continue

            # 处理标题
            if line.startswith('#'):
                level = line.count('#')
                text = line.lstrip('#').strip()
                if level == 1:
                    p = doc.add_heading(text, level=0)
                else:
                    p = doc.add_heading(text, level=min(level - 1, 4))
                # 强制覆盖运行级别字体
                for run in p.runs:
                    self._set_font(run.font, '微软雅黑')

            # 处理列表
            elif line.startswith('- ') or line.startswith('* '):
                text = line[2:].strip()
                p = doc.add_paragraph(text, style='List Bullet')
                for run in p.runs:
                    self._set_font(run.font, '微软雅黑', Pt(11))

            elif line.startswith('1. '):
                text = line[3:].strip()
                p = doc.add_paragraph(text, style='List Number')
                for run in p.runs:
                    self._set_font(run.font, '微软雅黑', Pt(11))

            # 处理表格（简化版）
            elif line.startswith('|') and '|' in line[1:]:
                cells = [cell.strip() for cell in line.split('|')[1:-1]]
                # 跳过分隔行
                if all(set(c) <= set('- :') for c in cells):
                    continue
                if len(cells) > 0:
                    table = doc.add_table(rows=1, cols=len(cells))
                    for i, cell in enumerate(cells):
                        table.cell(0, i).text = cell
                        for p in table.cell(0, i).paragraphs:
                            for run in p.runs:
                                self._set_font(run.font, '微软雅黑', Pt(11))

            # 处理普通段落
            elif line:
                p = doc.add_paragraph(line)
                for run in p.runs:
                    self._set_font(run.font, '微软雅黑', Pt(11))

        doc.save(str(output_path))
        return output_path
